

# imports
import os 
import xml.etree.ElementTree as ET
import xmltodict
import re
import glob 
import requests
import io
import json
import xml
import time
import pymysql as mdb
from docx import Document
from docx.shared import Inches


ENV_DIR_CODE = os.environ.get('DIR_CODE')
ENV_DIR_PUBMED = os.environ.get('DIR_PUBMED')
ENV_DIR_SUMMARY_DOC = os.environ.get('DIR_SUMMARY_DOC')

# import relative libraries
dir_code = "/home/javaprog/Code/PythonWorkspace/"
if ENV_DIR_CODE:
    dir_code = ENV_DIR_CODE
import sys
sys.path.insert(0, dir_code + 'MachineLearningPython/DccKP/GPT/')
import dcc_gpt_lib

# constants
DOC_FILENAME = "/home/javaprog/Data/Broad/GPT/GeneSummaries/geneSummary_{}_{}.docx"
if ENV_DIR_SUMMARY_DOC:
    DOC_FILENAME = ENV_DIR_SUMMARY_DOC + "/geneSummary_{}_{}_{}.docx"

DB_PASSWD = os.environ.get('DB_PASSWD')
SCHEMA_GPT = "gene_gpt"

# SQL_SELECT_GENE_SUMMARIES = """
# select se.gene, abst.abstract
# from pgpt_paper_abstract abst, pgpt_search se 
# where abst.search_top_level_of = se.id
# order by se.gene;
# """
SQL_SELECT_GENE_SUMMARY_BY_GENE = """
select se.gene, abst.abstract
from pgpt_paper_abstract abst, pgpt_search se 
where abst.search_top_level_of = se.id and abst.gpt_run_id = %s and se.gene = %s
"""

# methods
def get_db_gene_summaries(conn, id_run, list_genes=None, num_summaries=-1, log=False):
    '''
    get a list of abstract map objects
    '''
    # initialize
    list_summaries = []
    cursor = conn.cursor()

    # query 
    for gene in list_genes:
        if log:
            print("searching for gene: {}".format(gene))

        cursor.execute(SQL_SELECT_GENE_SUMMARY_BY_GENE, (id_run, gene))
        db_result = cursor.fetchall()
        for row in db_result:
            gene = row[0]
            abstract = row[1]
            list_summaries.append({"gene": gene, 'summary': abstract})

    # return
    return list_summaries



# main
if __name__ == "__main__":
    # get the db connection
    conn = dcc_gpt_lib.get_connection()
    list_genes_lipodystrophy = ['LMNA', 'PPARG', 'PLIN1', 'AGPAT2', 'BSCL2', 'CAV1', 'PTRF']
    list_genes_mody = ['GCK', 'HNF1A', 'HNF1B', 'CEL', 'PDX1', 'HNF4A', 'INS', 'NEUROD1', 'KLF11']
    list_genes = list_genes_mody

    # map
    map_gene_lists = {'T2D': dcc_gpt_lib.LIST_T2D, 'CAD': dcc_gpt_lib.LIST_CAD, 'Osteoarthritis': dcc_gpt_lib.LIST_OSTEO, 
                      'Kidney': dcc_gpt_lib.LIST_KCD, 'T1D': dcc_gpt_lib.LIST_T1D, 'Obesity': dcc_gpt_lib.LIST_OBESITY}
    
    map_runs = {'Genetics': 7, 'Biology': 8}

    for key_run, id_run in map_runs.items():
        # loop through map
        for key, value in map_gene_lists.items():
            # get the list of genes
            list_genes = dcc_gpt_lib.create_list_from_string_list([value])
            disease = key
            
            # list_summaries = get_db_gene_summaries(conn=conn)
            print("looking for run: {} - {} with disease: {}".format(key_run, id_run, disease))
            list_summaries = get_db_gene_summaries(conn=conn, id_run=id_run, list_genes=list_genes, log=True)
            print("found summary count: {} for run: {} with disease: {}".format(len(list_summaries), key_run, disease))

            if list_summaries and len(list_summaries) > 0:
                # create the document
                document = Document()

                # loop
                for item in list_summaries:
                    # get the data
                    gene = item.get('gene')
                    summary = item.get('summary')

                    # add to document
                    document.add_heading('Gene: {}'.format(gene), 0)

                    document.add_paragraph(summary, style='Intense Quote')

                    document.add_page_break()

                # save the document
                file_document = DOC_FILENAME.format(key_run, disease, round(time.time() * 1000))
                print("saving list to document: {}".format(file_document))
                document.save(file_document)



    #     document.add_heading('Document Title', 0)

    #     p = document.add_paragraph('A plain paragraph having some ')
    #     p.add_run('bold').bold = True
    #     p.add_run(' and some ')
    #     p.add_run('italic.').italic = True

    #     document.add_heading('Heading, level 1', level=1)
    #     document.add_paragraph('Intense quote', style='Intense Quote')

    #     document.add_paragraph(
    #         'first item in unordered list', style='List Bullet'
    #     )
    #     document.add_paragraph(
    #         'first item in ordered list', style='List Number'
    #     )

    #     document.add_picture('monty-truth.png', width=Inches(1.25))

    #     document.add_page_break()

    # document.save('demo.docx')

