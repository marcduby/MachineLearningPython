

# imports
import os 
import xml.etree.ElementTree as ET
import xmltodict
import re
import glob 
import requests
import io
import json
import xml
import time
import pymysql as mdb
from docx import Document
from docx.shared import Inches
from time import gmtime, strftime


# constants
ENV_DIR_CODE = os.environ.get('DIR_CODE')
ENV_DIR_PUBMED = os.environ.get('DIR_PUBMED')
ENV_DIR_SUMMARY_DOC = os.environ.get('DIR_SUMMARY_DOC')

# import relative libraries
dir_code = "/home/javaprog/Code/PythonWorkspace/"
if ENV_DIR_CODE:
    dir_code = ENV_DIR_CODE
import sys
sys.path.insert(0, dir_code + 'MachineLearningPython/DccKP/GPT/')
import dcc_gpt_lib

# constants
FILE_SUMMARY_DOC = "/home/javaprog/Data/ML/Llama2Test/PPARG/Docs/{}_{}_{}.docx"
DB_PASSWD = os.environ.get('DB_PASSWD')
SCHEMA_GPT = "gene_gpt"



# methods


# main
if __name__ == "__main__":
    # get the connection
    conn = dcc_gpt_lib.get_connection(schema=SCHEMA_GPT)

    # get the list of abstracts
    list_abstracts = dcc_gpt_lib.get_db_most_ref_abstracts_for_search(conn=conn, id_search=1, limit=300, to_shuffle=False)

    # loop and create
    if list_abstracts and len(list_abstracts) > 0:
        # loop
        for index, item in enumerate(list_abstracts):
            # create the document
            document = Document()

            # get the data
            title = item.get('title')
            pubmed_id = item.get('pubmed_id')
            abstract = item.get('abstract')
            ref_count = item.get('ref_count')

            # # add to document
            # document.add_heading('Gene: {}'.format(gene), 0)

            # document.add_paragraph(summary, style='Intense Quote')
            document.add_paragraph(abstract, style='Normal')

            # save the document
            str_time = strftime("%Y-%m-%d", gmtime())
            # file_document = DOC_FILENAME.format(key_run, disease, round(time.time() * 1000))
            file_document = FILE_SUMMARY_DOC.format(index, pubmed_id, ref_count)
            print("saving abstract: {} to document: {}".format(pubmed_id, file_document))
            document.save(file_document)

